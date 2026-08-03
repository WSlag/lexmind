"""Parsers that normalize an uploaded file into a :class:`Document`.

Supported formats: ``txt``, ``docx``, ``pdf``. Optional heavy dependencies
(``python-docx``, ``pdfplumber``) are imported lazily so the pipeline degrades
gracefully when they are missing.
"""
from __future__ import annotations

import re
from pathlib import Path

from app.schemas.review import Document

TXT_EXTENSIONS = {".txt", ".md", ".csv"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}

# Regex for likely section headings, e.g. "3.2 Pricing" or "4. Force Majeure".
_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)\s*[.)]?\s*([A-Za-z][A-Za-z \-&/]{2,})$")


class ParsingError(Exception):
    """Raised when a document cannot be parsed into text."""


def _ext(name: str) -> str:
    return Path(name).suffix.lower().lstrip(".")


def parse_file(filename: str, data: bytes) -> Document:
    """Parse raw file bytes into a normalized ``Document``."""
    ext = _ext(filename)
    dotted = f".{ext}"
    if dotted in TXT_EXTENSIONS:
        content = data.decode("utf-8", errors="replace")
    elif dotted in DOCX_EXTENSIONS:
        content = _parse_docx(data)
    elif dotted in PDF_EXTENSIONS:
        content = _parse_pdf(data)
    else:
        raise ParsingError(f"Unsupported file type: .{ext}")

    content = _normalize(content)
    words = content.split()
    return Document(
        filename=filename,
        file_type=ext if dotted in PARSEABLE else "txt",
        content=content,
        contract_type=_detect_contract_type(content),
        char_count=len(content),
        word_count=len(words),
    )


_CONTRACT_TYPE_HINTS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"construction agreement", re.I), "Commercial Construction Agreement"),
    (re.compile(r"gas supply agreement", re.I), "Gas Supply Agreement"),
    (re.compile(r"supply agreement", re.I), "Supply Agreement"),
    (re.compile(r"purchase agreement|sale agreement|master agreement", re.I), "Purchase / Sale Agreement"),
    (re.compile(r"service agreement|services agreement", re.I), "Service Agreement"),
    (re.compile(r"lease agreement|tenancy agreement", re.I), "Lease Agreement"),
    (re.compile(r"consulting agreement|consultancy", re.I), "Consulting Agreement"),
    (re.compile(r"licen[cs]e agreement", re.I), "Licence Agreement"),
    (re.compile(r"nda|non[- ]disclosure", re.I), "Non-Disclosure Agreement"),
    (re.compile(r"employment agreement|offer letter", re.I), "Employment Agreement"),
]


def _detect_contract_type(content: str) -> str:
    """Best-effort agreement-type detection from the document header text."""
    head = content[:2000]
    for pattern, label in _CONTRACT_TYPE_HINTS:
        if pattern.search(head):
            return label
    return "Agreement"


PARSEABLE = (TXT_EXTENSIONS | DOCX_EXTENSIONS | PDF_EXTENSIONS) - {".md"}


def _parse_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ParsingError("DOCX support requires 'python-docx'") from exc
    from io import BytesIO

    d = docx.Document(BytesIO(data))
    parts: list[str] = []
    for para in d.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise ParsingError("PDF support requires 'pdfplumber'") from exc
    from io import BytesIO

    with pdfplumber.open(BytesIO(data)) as pdf:
        return "\n\f".join(page.extract_text() or "" for page in pdf.pages)


def _normalize(text: str) -> str:
    """Collapse whitespace while preserving paragraph breaks."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()